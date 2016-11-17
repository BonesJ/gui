import os
import os.path
import sys
import sqlite3
import unicodedata as ucd
from PyQt4 import QtCore, QtGui, QtSql, uic  # QtWebKit,
from billQdialog import Ui_Dialog
from itertools import cycle
from functools import partial
import compiled_resources
import treeselector_
from datetime import datetime, date

# import treeviewmodel

# from kiwi_UI import Ui_MainWindow

# GLOBAL VAR
DB = "Alpha.db"

STYLES = ['blender_mod_btns.css', 'blank.css']
# , 'blender.css', 'dark-blue.css', 'dark-green.css','dark-orange.css',
# 'light-blue.css', 'light-green.css','light-orange.css']
STYLE = cycle(STYLES)
MODIF_LINE_COUNT = {"001": 5, "002": 3}

#  max line 2423


class MyWindow(QtGui.QMainWindow):
    """Main class in which the app runs.

    A bit hacky but it is a running application, so lets focus on getting
    most things working, then we will think about updating the whole
    structure so that it makes more sense.
    Notes and abbreviations:
    WF      Wrapper function
    """

    def __init__(self, model, model2, model3, parent=None):
        """Init taking all models of app. Add more arguments for more models.

        At present time we take in 3 models used for the tableviews
        """
        super(MyWindow, self).__init__(parent, QtCore.Qt.FramelessWindowHint |
            QtCore.Qt.WindowMinimizeButtonHint |
            QtCore.Qt.WindowMaximizeButtonHint)
        # self.setupUi(self)
    # ===========================================================FUNCTIONS DEFS

        def del_selected_jobs():
            """WF for jobs data using function `del_selected()`."""
            del_selected(self.jobsTable, model2, table="job")

        def del_selected_clients():
            """WF for clients data using function `del_selected()`."""
            del_selected(self.clientsTable, model)

        def combo_jobs_status():
            """Func to disable jobs btns unless sth selected in ComboBox."""
            # print self.comboBoxClients.currentText()
            if "<Select Client>" in self.comboBoxClients.currentText():
                self.btnNewJ.setEnabled(False)
                self.btnEditJ.setEnabled(False)
                self.btnDelJ.setEnabled(False)
            else:
                self.btnNewJ.setEnabled(True)
                self.btnEditJ.setEnabled(True)
                self.btnDelJ.setEnabled(True)

        def combo_emails_status():
            """Func to disable bills btns unless sth selected in ComboBox."""
            # print self.comboBoxEmails.currentText()
            if "<Custom>" in self.comboBoxEmails.currentText():
                self.leCustomEmail.setMinimumWidth(160)
                self.btnSendMail.setEnabled(False)
            elif "<Select Email>" in self.comboBoxEmails.currentText():
                self.btnSendMail.setEnabled(False)
                self.leCustomEmail.setMaximumWidth(0)
                self.leCustomEmail.setMinimumWidth(0)
            else:
                self.leCustomEmail.setMaximumWidth(0)
                self.leCustomEmail.setMinimumWidth(0)

        def update_combo_client_for_jobs():
            """Function update contents of ComboBox.

            Maybe needs to be rewritten to take following arguments:
            `ComboBox`, `ColumnName`, `Table` in order to fill ComboBox
            with rows from `ColumnName` in `Table`. To be addressed
            """
            # # UNCOMMENT SECTION FOR DEBUGGING /START/
            # print "\n"
            # print "-" * 80
            # print "START OF FUNCTION update combo client"
            # # UNCOMMENT SECTION FOR DEBUGGING /END/

            connection = sqlite3.connect(DB)
            cursor = connection.cursor()
            # # UNCOMMENT SECTION FOR DEBUGGING /START/
            # cursor.execute("""SELECT * FROM client""")
            # result_all = cursor.fetchall()
            # print "result all:", result_all
            # # UNCOMMENT SECTION FOR DEBUGGING /END/

            cursor.execute("""SELECT DISTINCT "Code Client" FROM client""")
            result_key = cursor.fetchall()
            connection.close()

            # UNCOMMENT SECTION FOR DEBUGGING /START/
            # print "result_key:", result_key

            # UNCOMMENT SECTION FOR DEBUGGING /END/

            for p in result_key:
                p = list(p)
                if p[0] is not None:
                    # # UNCOMMENT SECTION FOR DEBUGGING /START/
                    # print "unecoded", p
                    # # UNCOMMENT SECTION FOR DEBUGGING /END/

                    p = ucd.normalize('NFKD', p[0]).encode('ascii', 'ignore')
                    # # UNCOMMENT SECTION FOR DEBUGGING /START/
                    # print "encoded" , p
                    # # UNCOMMENT SECTION FOR DEBUGGING /END/
                else:
                    # # UNCOMMENT SECTION FOR DEBUGGING /START/
                    # print p, "DID NOT ENCODE"
                    # # UNCOMMENT SECTION FOR DEBUGGING /END/
                    return
                for i in range(self.comboBoxClients.count()):

                    # # retrieve combox current items
                    allitems = [str(self.comboBoxClients.itemText(u))
                                for u in range(self.comboBoxClients.count())]
                    # # UNCOMMENT SECTION FOR DEBUGGING /START/
                    # print "All items:", allitems
                    # # UNCOMMENT SECTION FOR DEBUGGING /END/

                    if p not in allitems:
                        # # UNCOMMENT SECTION FOR DEBUGGING /START/
                        # print "DID NOT ADD ITEM TO COMBO BOX"
                        # # add the name of the client by list compreh. search
                        # tup = [item for item in result_all if p in item]
                        # print "code:", tup
                        # # UNCOMMENT SECTION FOR DEBUGGING /END/
                        self.comboBoxClients.addItem(p)
                        # # UNCOMMENT SECTION FOR DEBUGGING /START/
                        # print "ADDED ITEM %s TO COMBO BOX" % p
                        # print "All items:", allitems
                        # # UNCOMMENT SECTION FOR DEBUGGING /END/

            # # UNCOMMENT SECTION FOR DEBUGGING /START/
            # print "END OF FUNCTION update combo client"
            # print "-" * 80
            # print "\n"
            # # UNCOMMENT SECTION FOR DEBUGGING /END/

        def job_bill_btnstate(b, combo, group=None):

            # print "button: ", b.text()
            combo.setEnabled(True)
            combo.setMinimumWidth(150)

            if b.text() == "Clients":
                # combo.setEnabled(True)
                # combo.setMinimumWidth(150)
                combo.addItem("<Select Client>")
                if b.isChecked():
                    # print b.text() + " is selected"
                    update_combo_from_db(combo=combo,
                                         column="Code Client", table="client")

                    args = partial(update_combo_from_db, combo=combo,
                                   column="Code Client", table="client")
                    combo.highlighted.connect(args)
                else:
                    # print b.text() + " is deselected"
                    combo.highlighted.disconnect()
                    combo.clear()

            if b.text() == "Jobs":
                # combo.setEnabled(True)
                # combo.setMinimumWidth(150)
                combo.addItem("<Select Job>")

                if b.isChecked():
                    # print b.text() + " is selected"
                    update_combo_from_db(combo=combo,
                                         column="Code Job",
                                         table="job")
                    args = partial(update_combo_from_db,
                                   combo=combo,
                                   column="Code Job",
                                   table="job")
                    combo.highlighted.connect(args)
                else:
                    # print b.text() + " is deselected"
                    combo.highlighted.disconnect()
                    combo.clear()

            if b.text() == "All":
                combo.setMinimumWidth(0)
                combo.setMaximumWidth(0)
                combo.setEnabled(False)
                if group == "jobs":
                    get_jobs_tabview(self.comboBoxClients, self.jobsTable,
                             self.model2, showall=True)
                elif group == "bills":
                    get_bills(self.comboBoxBills, self.billsTable,
                             self.model3, "confirmed", showall=True)
                # print "Will print all and disable combobox"

        def update_combo_client_for_bills():
            update_combo_from_db(self.comboBoxBills, "Code Client", "client")

        def get_jobs_tabview_wrapper():
            """WF for jobs using function `get_jobs_tabview()`.

            Passed args are:
            self.comboBoxClients, self.jobsTable, self.model2
            """
            get_jobs_tabview(self.comboBoxClients, self.jobsTable, self.model2)

        def dock_toggle():
            """Func to toggle `Modifications` dock."""
            if self.dockWidget.isVisible() is True:
                self.dockWidget.close()
            else:
                self.dockWidget.show()

        def change_theme():
            """Func to change loaded Qt StyleSheet from menu."""
            global STYLE
            css = QtCore.QFile(os.path.join('themes', STYLE.next()))
            css.open(QtCore.QIODevice.ReadOnly)

            if css.isOpen():
                self.setStyleSheet(QtCore.QVariant(css.readAll()).toString())
            css.close()

        def close_application():
            QtGui.QApplication.exit()
            print 'Application closed successfully'


        def new_client():
            """WF for `input_window()` to create new client.

            Passed args are:
            tab='client', fields=info2grab, mode='new',
            model=model, tabview=self.clientsTable
            """
            cols_i = [0, 1, 2, 3, 4, 5, 6, 8]

            info2grab = ["Organisation", "Nom", "Prenom", "Tel Mobile",
            "Tel Fixe", "Email"]
            input_window(tab='client', fields=info2grab,
                         mode='new', model=model,
                         tabview=self.clientsTable, cols_to_fetch=cols_i)

        def edit_client():
            """WF for `input_window()` to edit new client.

            Passed args are:
            tab='client', fields=info2grab, mode='edit',
            model=model, tabview=self.clientsTable
            """
            # cols index in tableview to fetch, 0=key, 2="Nom" etc...
            cols_i = [0, 1, 2, 3, 4, 5, 6, 8]
            # button names to create = column names in sqlite3
            info2grab = ["Organisation", "Nom", "Prenom", "Tel Mobile",
            "Tel Fixe", "Email"]
            input_window(tab='client', fields=info2grab,
                         mode='edit', model=model,
                         tabview=self.clientsTable, cols_to_fetch=cols_i)

        def new_job():
            """WF for `input_window()` to create new job.

            Passed args are:
            tab='client', fields=info2grab, mode='new',
            model=model, tabview=self.clientsTable
            """
            cols_i = [0, 1, 2, 3, 4, 5, 6, 7]

            info2grab = ["Code Job", "Produit", "Brief", "Date Debut",
                         "Date Fin", "Prix", "Somme Payee"]
            input_window(tab='job', fields=info2grab,
                         mode='new', model=model2, combo=self.comboBoxClients,
                         tabview=self.jobsTable, cols_to_fetch=cols_i)

        def edit_job():
            # cols_i = [0, 1, 2, 3, 4, 5, 6]
            cols_i = [0, 1, 2, 3, 4, 5, 6, 7]
            info2grab = ["Code Job", "Produit", "Brief", "Date Debut",
                         "Date Fin", "Prix", "Somme Payee"]
            input_window(tab='job', fields=info2grab,
                         mode='edit', model=model2,
                         tabview=self.jobsTable, cols_to_fetch=cols_i)

        def new_bill_input_wrapper():
            """WF for bills using function `new_bill_input_dialog()`.

            Passed args are:
            self.comboBoxClients, self.billsTable
            """
            new_bill_input_dialog(combo=self.comboBoxClients,
                                  table=self.billsTable,
                                  model=model4)

        def getfiles(self):
            """Func to display window to retrieve file."""
            dlg = QtGui.QFileDialog()
            dlg.setFileMode(QtGui.QFileDialog.AnyFile)
            dlg.setFilter("SQLite file (*.db)")

            set_style(dlg)
            filenames = QtCore.QStringList()

            if dlg.exec_():

                filenames = dlg.selectedFiles()
                f = open(filenames[0], 'r')

                with f:
                    data = f.read()
                    self.contents.setText(data)

        def fetch_jobs_from_client():
            forclient = str(self.modifClientComboBox.currentText())
            # print "slected client :>:>", forclient
            if forclient == "<Client>":
                self.modifJobComboBox.setEnabled(False)
            else:
                self.modifJobComboBox.setEnabled(True)
                # print "passing>> %s in displayfor" % forclient


                update_combo_from_db(
                    combo=self.modifJobComboBox,
                    column="Code Client", table="job",
                    displayfor=forclient)

                # args = partial(update_combo_from_db,
                #     combo=self.modifJobComboBox,
                #     column="Code Client", table="job",
                #     displayfor=forclient)

                # self.modifJobComboBox.highlighted.connect(args)
        def style_modif_close_btn(button):
            button.setStyleSheet("""
            QPushButton {
                color: #FFFFFF;
                font: 9pt bold "Roboto" ;
                text-align: center;
                background-color: rgba(0,0,0,0);

                border: 1px solid black;
                border-radius: 2px;
                min-height: 0px; /* same as QTabBar QPushButton min-width */
                min-width: 0px;
                width: 10px;
                height: 12px;
            }


            QPushButton:hover {
                /*border: 1px solid white;*/
                background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0, y2:1, stop:0 #1565C0, stop:1 #1565C0);
                background-color: rgba(255,255,255,50);
                background-color:rgb(100, 0, 0);
                border-style: inset;
                border-color:rgb(100, 0, 0);


            }

            QPushButton:disabled,
            QPushButton:disabled:checked {
                color: #4c4c4c;
                background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0, y2:1, stop:0 #878787, stop:1 #7c7c7c);
                border-color: #5b5b5b;
            }

            QPushButton:pressed {
                color: white;
                background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0, y2:1, stop:0 #1565C0, stop:1 #0D47A1);
                background-color:rgb(200, 0, 0);
            }
            """)

        def add_modif_line():
            # insert row of checkbox, lineedit, and 'x' close
            # insert row into debrief table with timestamp and ischecked? status

            def strikeout(self):
                if checkboxes["001"].isChecked():
                    lines["001"].setStyleSheet("""
                       QLineEdit
                       {
                            text-decoration: line-through;
                       }
                       """)
                else:
                    lines["001"].setStyleSheet("""
                       QLineEdit
                       {

                       }
                       """)

            def remove_line(self):
                pass

            global MODIF_LINE_COUNT

            # input_map = dict(zip(fields, cols_to_fetch))
            # print "input map", input_map

            # Fetch all modifs for current job
            #   create rows in dockedwidget
            #   check status of modif
            # Update records

        def del_modif(num=None,btn=None, line=None, close=None, modif=None):
            """Modification dock function to delete rows.

            Deletes a row of a given modification for a given
            job (or Projet). Clears databse table 'debrief' of rows.

            Args:
                btn  (QCheckBox):  Delets widget.
                line (QLineEdit):  Delets widget.
                modif (tup ):      Read information (id and contents)
            Returns:
                True
            """
            #  Warning dialog
            msg = QtGui.QMessageBox()
            msg.setIcon(QtGui.QMessageBox.Warning)
            msg.setText("Etes vous sur de vouloir supprimer cette modification ? Celle-ci sera definitivement efface de la memoire. Appuiyer sur 'Ok' pour continuer ou sur 'Cancel' pour annuler.")
            # msg.setInformativeText("To edit a client you must select a row.")
            msg.setWindowTitle("Delete")
            msg.setStandardButtons(QtGui.QMessageBox.Ok | QtGui.QMessageBox.Cancel)
            set_style(msg)
            result = msg.exec_()

            if result == QtGui.QMessageBox.Ok:
                # Remove widget
                widgets = [num, btn, line, close]
                for widget in widgets:
                    print "Removed", widget
                    self.gridLayout_21.removeWidget(widget)
                    widget.deleteLater()
                    widget = None

                # Remove recod from db
                connection = sqlite3.connect(DB)
                cursor = connection.cursor()
                cursor.execute("""
                        DELETE FROM debrief
                        WHERE "key" = "{id}" """.format(id=modif[0]))

                result_all = cursor.fetchall()
                connection.commit()
                connection.close()

        def fetch_modif_items(parent=None):
            """Modification dock function for displaying edit boxes.

            Fetches number of modifs stored in debrief table and for r rows pointed to
            `parent`, it generates r rows of a QCheckBox, QLineEdit and QPushButton.
            -At generation, it also checks for [Date_Done] in table debrief, and if
            date exists, then set checkbox as checked and setstyle of lineedit
            to strike-out.

            Args:
                parent (str): "Project_Code" from modifComboBox.
            Returns:
                True
            """
            def update_progressbar():
                # Add progress barr percentage
                checked_items = 0
                all_items = 0

                for item in checkboxes:
                    all_items += 1
                    if checkboxes[item].isChecked():
                        checked_items += 1

                # print "checked_items", checked_items, "/", all_items
                if all_items == 0:
                    self.progressBar.setValue(0)
                else:
                    value = float(checked_items) / float(all_items) * 100
                    print "Modifications are %s%% complete" % (int(value))
                    self.progressBar.setValue(int(value))

                # Add progress barr percentage

            def checkbox_state(btn=None, line=None, modif=None):
                """Toggle done editing modif.

                Toggles the QLineedit strikeout text. If striked out
                then created time stamp for DTAE DONE in debrief table
                ad insert value and Disable editing in Qlineedit.

                Args:
                    btn  (QCheckBox):  Reads state of check box.
                    line (QLineEdit):  Sets text style as strikeout and
                                       disables editing
                    modif (tup ):      Read information (id and contents)
                Returns:
                    True
                """
                if btn.isChecked() is True:
                    line.setStyleSheet("""
                       QLineEdit
                       {
                            text-decoration: line-through;
                            background-color: rgba(0,255,0,85);
                            color: white;
                       }
                       """)
                    import datetime
                    now = str(datetime.datetime.now())
                    # print('Date now: %s' % now)
                    # now = now.replace(":", "-")
                    # now = now.replace(" ", "-")
                    # now = now.replace(".", "-")
                    print('DateDone replaced: %s' % now)
                    date_completed = now

                    # SQlite code to update date done to now
                    connection = sqlite3.connect(DB)
                    cursor = connection.cursor()
                    cursor.execute("""
                            UPDATE debrief SET "Date_Done" = "{date}"
                            WHERE "key" = "{id}" """
                                .format(id=modif[0], date=date_completed))
                    connection.commit()
                    connection.close()

                    # insert code to timestamp strike trough even above
                    line.setReadOnly(True)

                    update_progressbar()
                else:
                    line.setStyleSheet("""
                       QLineEdit
                       {
                            text-decoration: none;


                       }
                       """)

                    # SQlite code to update date done to NULL
                    date_completed = "NULL"

                    connection = sqlite3.connect(DB)
                    cursor = connection.cursor()
                    cursor.execute("""
                            UPDATE debrief SET "Date_Done" = "{date}"
                            WHERE "key" = "{par_id}" """
                                .format(par_id=modif[0], date=date_completed))
                    connection.commit()
                    connection.close()
                    print('DateDone replaced: %s' % date_completed)

                    line.setReadOnly(False)
                    update_progressbar()

            # clear layout of previous widgets
            widgets = self.frame_7.children()
            # print "WIDGETS", widgets
            for widget in widgets:
                # print widget.objectName(), "is of type: ", type(widget)
                if type(widget) == type(self.gridLayout_21):
                    print
                    print "Found a layout, will not remove !"
                    print
                    continue
                else:
                    self.gridLayout_21.removeWidget(widget)
                    widget.deleteLater()
                    widget = None

            # lookup debrief table and select chosen job to display modif
            print
            print "Fetching modifs for job: ",
            modif_todo_job_id = str(self.modifJobComboBox.currentText())
            parent = modif_todo_job_id[0:9]
            print parent

            if parent is None:
                print 'No parent specified for fetch_modif_line()'

            else:
                # fetch records for select job
                connection = sqlite3.connect(DB)
                cursor = connection.cursor()
                cursor.execute("""
                        SELECT * FROM debrief
                        WHERE "Parent" = "{par_id}" """.format(par_id=parent))

                result_all = cursor.fetchall()
                connection.commit()
                connection.close()
                # print result_all

                # extract modifications count from result_all
                modif_line_count = len(result_all)
                print "Project ", "has ", modif_line_count, "modifs."

                i = 0
                closebuttons = {}
                lines = {}
                checkboxes = {}
                row_count = 1

                # Row factory
                for modif in result_all:
                    print modif
                    id = modif[0]
                    desc = modif[1]

                    print id, desc

                    numbering = QtGui.QLabel(str(row_count))
                    checkboxes[id] = QtGui.QCheckBox()

                    lines[id] = QtGui.QLineEdit(desc)
                    closebuttons[id] = QtGui.QPushButton('X')
                    style_modif_close_btn(closebuttons[id])

                    self.gridLayout_21.addWidget(numbering, i, 0, 1, 1)
                    self.gridLayout_21.addWidget(checkboxes[id], i, 1, 1, 1)
                    self.gridLayout_21.addWidget(lines[id], i, 2, 1, 2)
                    self.gridLayout_21.addWidget(closebuttons[id], i, 3, 1, 1)

                    #  Fix widget and prevent it from stretching
                    sizePolicy = QtGui.QSizePolicy(
                        QtGui.QSizePolicy.Fixed, QtGui.QSizePolicy.Fixed)
                    sizePolicy.setHeightForWidth(True)
                    numbering.setSizePolicy(sizePolicy)
                    checkboxes[id].setSizePolicy(sizePolicy)
                    closebuttons[id].setSizePolicy(sizePolicy)

                    # connect checkbox signals
                    checkboxes[id].toggled.connect(partial(checkbox_state,
                        btn=checkboxes[id], line=lines[id], modif=modif))

                    closebuttons[id].clicked.connect(partial(del_modif,
                        num=numbering,btn=checkboxes[id], line=lines[id],
                        close=closebuttons[id], modif=modif))


                    # check db for "Date Done", if null leave checkbox unchecked,
                    # else make checkbox checked .setChecked(checked)
                    if modif[3] == "NULL":
                        checkboxes[id].setChecked(False)
                    else:
                        checkboxes[id].setChecked(True)

                    # checkboxes[id].deleteLater()
                    # lines[id].deleteLater()
                    # closebuttons[id].deleteLater()

                    i += 1
                    row_count += 1

                update_progressbar()

        def save_modif():
            modif_todo_job_id = str(self.modifJobComboBox.currentText())
            parent = modif_todo_job_id[0:9]

            # fetch records for select job
            connection = sqlite3.connect(DB)
            cursor = connection.cursor()
            cursor.execute("""
                    SELECT * FROM debrief
                    WHERE "Parent" = "{par_id}" """.format(par_id=parent))

            result_all = cursor.fetchall()
            connection.commit()
            connection.close()

            for modif in result_all:
                print modif

        def add_modif_item():
            """Modification dock function for displaying adding boxes.

            Addes a modification row for given parent (selected job)

            Args:
                parent (str): "Project_Code" from modifComboBox.
            Returns:
                None
            """
            modif_todo_job_id = str(self.modifJobComboBox.currentText())
            parent = modif_todo_job_id[0:9]
            print parent
            if parent is None:
                print 'No parent specified for add_modif_line()'
            else:
                import datetime

                now = str(datetime.datetime.now())
                print('Date now: %s' % now)
                # now = now.replace(":", "-")
                # now = now.replace(" ", "-")
                # now = now.replace(".", "-")
                # print('Date now replaced: %s' % now)

                description = 'Entrez votre debrief'
                date_initialised = now
                date_completed = 'NULL'

                debrief_data = [(description, date_initialised, date_completed,
                        parent)]
                connection = sqlite3.connect(DB)
                cursor = connection.cursor()
                cursor.executemany("""
                    INSERT INTO debrief ("key","Description", "Date_Created",
                                               "Date_Done", "Parent")
                    VALUES(NULL,?,?,?,?)""", debrief_data)
                connection.commit()
                connection.close()

                fetch_modif_items(parent)

        def clients_sel_changed():
                data = tabView_sel_changed(tabview=self.clientsTable)

                # index = self.clientsTable.selectedIndexes()[0]
                # print('Row %d is selected' % index.row())
                # list_of_indices = self.clientsTable.selectedIndexes()
                # # print "Lst of indices:", list_of_indices
                # o = 0
                # for index in sorted(list_of_indices):
                #     print "Column ", index.column()
                #     id_us = str(self.clientsTable.model().data(index).toString())
                #     print "Data :", id_us
                print data
                self.label_55.setText(data[1])
                self.label_56.setText(data[2]+" "+data[3]+" "+data[4])
                # set below to labels of facturation total, facturation Payee etc

                # self.label_55.setText(data[1])
                # self.label_55.setText(data[1])
                # self.label_55.setText(data[1])

        def save_cloud():
            self.statusBar.showMessage("Telechargement vers un serveur, veuillez patienter...", 0)
            import upload
            upload.main()
            self.statusBar.showMessage("Done uploading file", 4000)

        def save_local():
            # Implement a save local database
            pass

        def minimise_main_window():
            self.showMinimized()

    # ===============================================APPLICATION INITIALISATION

        # In here configure applicationwide settings(theme, dock, menubar etc.)

        # # load ui
        uic.loadUi('kiwi_UI.ui', self)
        change_theme()

        label_copyright = QtGui.QLabel('Copyright P2 Labs. Algeria 2016')

        label_copyright.setStyleSheet(
            "QLabel { "
            "background-color: rgba(255, 255, 255, 0);"
            "border-style: inset; border-width: 1px; "
            "border-radius: 10px; border-color: beige; "
            "font: bold 8px;}")
        self.statusBar.addPermanentWidget(label_copyright)

        # close dock widget by default
        self.dockWidget.close()
        self.btnModif.clicked.connect(dock_toggle)
        self.btnSaveCloud.clicked.connect(save_cloud)
        self.btnSaveLocal.clicked.connect(save_local)

        # menubar bar
        self.actionCycle_Theme.triggered.connect(change_theme)
        self.actionLoad.triggered.connect(getfiles)
        self.actionQuitter_Application.triggered.connect(close_application)
        self.actionMinimiser_la_Fenetre.triggered.connect(minimise_main_window)
    # ========================================================MODIFICATION DOCK
        self.btnAddModif.clicked.connect(add_modif_line)

        args = partial(update_combo_from_db,
            combo=self.modifClientComboBox,
            column="Code Client",
            table="client")
        self.modifClientComboBox.highlighted.connect(args)
        self.modifClientComboBox.currentIndexChanged.connect(
            fetch_jobs_from_client)

        self.modifJobComboBox.activated.connect(fetch_modif_items)
        self.btnSaveModifs.clicked.connect(save_modif)
        self.btnAddModif.clicked.connect(add_modif_item)

        # self.modifJobComboBox.currentIndexChanged.connect(fetch_modif_todo)

        # Service proposes TAB, uses the edit class of ui widget
        tree_selector = init_tree_view(mode='edit')
        self.gridLayout_19.addWidget(tree_selector, 1, 0, 1, 2)

    # =============================================================FACTURES TAB

        self.model = model  # client table
        self.model2 = model2  # job table
        self.model3 = model3  # bill table
        self.model4 = model4  # invoice_items table

        # # disable  Bills and Jobs buttons as Default
        self.btnNewJ.setEnabled(False)
        self.btnEditJ.setEnabled(False)
        self.btnDelJ.setEnabled(False)
        self.btnNewBill.setEnabled(True)

        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
        # Clients groupbox Managment
        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
        # ========================================================Buttons Setup
        self.btnNewC.clicked.connect(new_client)
        self.btnEditC.clicked.connect(edit_client)
        self.btnDelC.clicked.connect(del_selected_clients)

        # ==========================================================Table Setup
        self.clientsTable.setModel(self.model)  # load db into tableClients
        self.clientsTable.setEditTriggers(
            QtGui.QAbstractItemView.NoEditTriggers)
        # trick to auto resize columns, hide, configure then show.
        self.clientsTable.setVisible(False)
        self.clientsTable.resizeColumnsToContents()
        self.clientsTable.resizeRowsToContents()
        hh = self.clientsTable.horizontalHeader()
        hh.setStretchLastSection(True)
        self.clientsTable.setShowGrid(False)
        self.clientsTable.setSortingEnabled(True)
        self.clientsTable.setVisible(True)

        selectionModel = self.clientsTable.selectionModel()
        selectionModel.selectionChanged.connect(clients_sel_changed)

        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
        # Jobs groupbox Managment
        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
        # ========================================================Buttons Setup
        # Blue Push buttons
        self.btnNewJ.clicked.connect(new_job)
        self.btnEditJ.clicked.connect(edit_job)
        self.btnDelJ.clicked.connect(del_selected_jobs)

        # Radio buttons
        JClients = partial(job_bill_btnstate, b=self.radioButtonJClients,
                                              combo=self.comboBoxClients,
                                              group="jobs")
        JAll = partial(job_bill_btnstate, b=self.radioButtonBAll,
                                          combo=self.comboBoxClients,
                                          group="jobs")

        self.radioButtonJClients.toggled.connect(JClients)  # Args too long...
        self.radioButtonJAll.toggled.connect(JAll)

        # View selection ComboBox
        self.comboBoxClients.setMinimumWidth(0)  # initially hide ComboBoxe
        self.comboBoxClients.setMaximumWidth(0)
        self.comboBoxClients.currentIndexChanged.connect(combo_jobs_status)
        self.comboBoxClients.activated.connect(get_jobs_tabview_wrapper)

        # ==========================================================Table Setup
        #  Enter code here...

        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
        # Bills groupbox Managment
        # ++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++
        # ========================================================Buttons Setup
        # Blue Push buttons
        self.btnNewBill.clicked.connect(new_bill_input_wrapper)

        # Radio buttons
        BClients = partial(job_bill_btnstate, b=self.radioButtonBClients,
                                              combo=self.comboBoxBills,
                                              group="bills")

        # BJobs = partial(job_bill_btnstate, b=self.radioButtonBJobs,
        #                                    combo=self.comboBoxBills,
        #                                    group="bills")

        BAll = partial(job_bill_btnstate, b=self.radioButtonBAll,
                                          combo=self.comboBoxBills,
                                          group="bills")

        self.radioButtonBClients.toggled.connect(BClients)  # Args too long...
        # self.radioButtonBJobs.toggled.connect(BJobs)
        self.radioButtonBAll.toggled.connect(BAll)

        # View selection ComboBox
        self.comboBoxBills.setMinimumWidth(0)  # initially hide ComboBoxe
        self.comboBoxBills.setMaximumWidth(0)
        # URGENT BELOW
        args = partial(get_bills, self.comboBoxBills, self.billsTable,
                       self.model3, "confirmed")
        self.comboBoxBills.activated.connect(args)

        # self.comboBoxEmails.currentIndexChanged.connect()  # NYI, soon will
        # self.comboBoxEmails.highlighted.connect()  # NYI, soon will

        self.leCustomEmail.setMinimumWidth(0)  # initially hide LineEdit
        self.leCustomEmail.setMaximumWidth(0)
        self.comboBoxEmails.activated.connect(combo_emails_status)

        # ==========================================================Table Setup
        # Enter code here

    # =============================================================CHARGES TAB
    # Add code here...
    # =============================================================PANORAMA TAB
    # Add code here
    # =============================================================DEVIS TAB

        self.show()

    def mousePressEvent(self, event):
        self.offset = event.pos()

    def mouseMoveEvent(self, event):
        x = event.globalX()
        y = event.globalY()
        x_w = self.offset.x()
        y_w = self.offset.y()
        self.move(x - x_w, y - y_w)

class MyModel(QtSql.QSqlTableModel):
    """Load model for DB manipulation using forms and QTableview."""

    def __init__(self, parent=None, tab=None):
        """Instantiate `QtSql.QSqlTableModel` & select sqlite3 table `tab`.

        Args:
            tab (QtGui.QTableview): The table to SELECT using sql

        Returns:
            bool: The return value. True for success, False otherwise.
        """
        super(MyModel, self).__init__(parent)
        self.tab = tab
        self.setEditStrategy(QtSql.QSqlTableModel.OnFieldChange)
        self.setTable(tab)
        self.select()

    def data(self, index, role=QtCore.Qt.DisplayRole):
        """Center align data in QTableview."""
        # self.tab = tab
        # print "table:", tab
        # cols_to_center_align
        # c_t_c_a = [1, 5, 6]
        # index.column() in c_t_c_a and
        if (role == QtCore.Qt.TextAlignmentRole):
            return QtCore.Qt.AlignCenter
        return QtSql.QSqlTableModel.data(self, index, role)


def set_style(widget):
    """Func used to apply css styling to window. Uses global variables."""
    global STYLES
    global STYLE
    for i in range((len(STYLES) - 1)):
        STYLE.next()

    css = QtCore.QFile(os.path.join('themes', STYLE.next()))
    css.open(QtCore.QIODevice.ReadOnly)
    if css.isOpen():
        widget.setStyleSheet(QtCore.QVariant(css.readAll()).toString())
    else:
        print "set_style() failed !"
    css.close()


def make_db():
    """Func used to create new SQlite DB for app usage.

    Should remove hardcoded data when done...
    """
    connection = sqlite3.connect(DB)
    cursor = connection.cursor()
    # DB has three main tables and 2 linking tables.
    # Code copied and pasted 3 times... need to find sol, in te mean time it
    # works...

    # clients table ...
    cursor.execute("""
        CREATE TABLE client (
        "key" INTEGER PRIMARY KEY AUTOINCREMENT,
        "Code Client" TEXT,
        "Organisation" TEXT,
        "Nom" TEXT,
        "Prenom" TEXT,
        "Tel Mobile" TEXT,
        "Tel Fixe" TEXT,
        "Factures Payees" TEXT,
        "Email" TEXT,
        "Solde Compte" TEXT,
        UNIQUE ("key") );""")

    # Populate db with some data ...
    cursor = connection.cursor()
    staff_data = [
        (55, "kaloumba", "James", "Bond", "05 52 44 04 77", "021 69 74 49", "5/7", "ceemss.saiatgmail", "DA 150 000.00" ),
        (55, "sadoomba", "Doctor", "Pepper", "05 52 44 04 77", "021 69 74 49", "5/7", "ceemss.saiatgmail", "DA 150 000.00" ),
        (55, "badoomba", "Coca", "Cola", "05 52 44 04 77", "021 69 74 49", "5/7", "ceemss.saiatgmail", "DA 150 000.00" ),
        (55, "frambounda", "Arya", "Stark", "05 52 44 04 77", "021 69 74 49", "5/7", "ceemss.saiatgmail", "DA 150 000.00" ),
    ]
    cursor.executemany("""
        INSERT INTO client ("key","Code Client", "Organisation", "Nom", "Prenom", "Tel Mobile",
            "Tel Fixe", "Factures Payees", "Email", "Solde Compte")
        VALUES(NULL,?,?,?,?,?,?,?,?,?)""", staff_data)
    connection.commit()

    # Fetch entire table of clients and store in `result_all`, and fetch
    # key data from first colmun `i[0]` in result_all through list compreh.
    # and store that  in result_cc (client code).
    # Could have used `SELECT key from client` but needded info for DEBUGGING
    #
    cursor.execute("""SELECT * FROM client""")
    result_all = cursor.fetchall()
    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    # print "-" * 80
    # print "-" * 80
    # print("fetchall before cc update:")
    # for r in result_all:
    #     print(r)
    # # UNCOMMENT SECTION FOR DEBUGGING /END/

    result_cc = [i[0] for i in result_all]
    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    # print "result_cc = ", result_cc
    # # UNCOMMENT SECTION FOR DEBUGGING /END/

    # Code below takes the unique key generated by SQL and appends `00` to it
    # to make it nicer looking when printing client numbers
    for i in result_cc:
        # # UNCOMMENT SECTION FOR DEBUGGING /START/
        # print "i=", i
        # # UNCOMMENT SECTION FOR DEBUGGING /END/
        count = "00" + str(i)
        count = count[-3:]
        # # UNCOMMENT SECTION FOR DEBUGGING /START/
        # print "concatenated:", count
        # # UNCOMMENT SECTION FOR DEBUGGING /END/
        cursor.execute("""
            UPDATE client SET "Code Client" = SUBSTR('000'||%s,-3, 3)
            WHERE "key" = %d  """ % (count, i))
    connection.commit()

    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    # cursor = connection.cursor()
    # cursor.execute("SELECT * FROM client")
    # result_all = cursor.fetchall()
    # print "-" * 80
    # print("fetchall after cc update:")
    # for r in result_all:
    #     print(r)
    # print "-" * 80
    # print "-" * 80
    # # UNCOMMENT SECTION FOR DEBUGGING /END/

    # Jobs table ...
    cursor.execute("""
        CREATE TABLE job (
        "key" INTEGER PRIMARY KEY ,
        "Code Job" VARCHAR(9),
        "Produit" TEXT,
        "Brief" VARCHAR(30),
        "Date Debut" VARCHAR(10),
        "Date Fin" VARCHAR(10),
        "Prix" VARCHAR(30),
        "Somme Payee" VARCHAR(30),
        UNIQUE ("key"));""")

    job_data = [
        ("J/001/001", "client 1", "descri", "15/06/2015", "17/10/2017", 15, 5),
        ("J/001/002", "client 1", "descri", "15/06/2015", "17/10/2017", 15, 5),
        ("J/002/003", "client 2", "descri", "15/06/2015", "17/10/2017", 15, 5),
        ("J/002/004", "client 2", "descri", "15/06/2015", "17/10/2017", 15, 5),
        ("J/002/005", "client 2", "descri", "15/06/2015", "17/10/2017", 15, 5),
        ("J/003/006", "client 3", "descri", "15/06/2015", "17/10/2017", 15, 5),
        ("J/003/007", "client 3", "descri", "15/06/2015", "17/10/2017", 15, 5),
    ]
    cursor.executemany("""
        INSERT INTO job ("key", "Code Job","Produit", "Brief", "Date Debut",
                         "Date Fin", "Prix", "Somme Payee")
        VALUES(NULL,?,?,?,?,?,?,?)""", job_data)
    connection.commit()

    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    # cursor.execute("SELECT * FROM job")
    # print("fetchall jobs:")
    # result = cursor.fetchall()
    # for r in result:
    #     print(r)
    # # UNCOMMENT SECTION FOR DEBUGGING /END/

    # "orders" table linking jobs to parent clients
    cursor.execute("""
        CREATE TABLE orders (
        key INTEGER PRIMARY KEY,
        client_id INTEGER,
        job_id INTEGER,
        FOREIGN KEY(client_id) REFERENCES client(key),
        FOREIGN KEY(job_id) REFERENCES job(key)
        );""")

    order_data = [(1, 1),
                  (1, 2),
                  (2, 3),
                  (2, 4),
                  (2, 5),
                  (3, 6),
                  (3, 7),
                  ]
    cursor.executemany("""
        INSERT INTO orders ("key","client_id", "job_id")
        VALUES(NULL,?,?)""", order_data)
    connection.commit()

    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    # cursor.execute("SELECT * FROM orders")
    # result_all = cursor.fetchall()
    # print "-" * 80
    # print("fetchall orders:")
    # for r in result_all:
    #     print(r)
    # print "-" * 80
    # # UNCOMMENT SECTION FOR DEBUGGING /END/

    # Bills table ...
    cursor.execute("""
        CREATE TABLE bill (
        "key" INTEGER PRIMARY KEY AUTOINCREMENT,
        "Code Facture" VARCHAR(10),
        "Produit" TEXT,
        "NAP" VARCHAR(30),
        "Date Facture" VARCHAR(30),
        "Date Paiement" VARCHAR(30),
        "Pourcentage Paye" REAL,
        "Statut" VARCHAR(30),
        UNIQUE ("key","Code Facture")
        );""")

    bill_data = [
        ("FAC001/001", "bill description...", 15000,
            "15/06/2015", "17/10/2017", 50, "Publie"),
        ("FAC002/002", "bill description...", 15000,
            "15/06/2015", "17/10/2017", 50, "Publie"),
        ("FAC002/003", "bill description...", 15000,
            "15/06/2015", "17/10/2017", 50, "Publie"),
        ("FAC003/004", "bill description...", 15000,
            "15/06/2015", "17/10/2017", 50, "Publie"),
        ("FAC003/005", "bill description...", 15000,
            "15/06/2015", "17/10/2017", 50, "Publie"),
    ]
    cursor.executemany("""
        INSERT INTO bill ("key","Code Facture", "Produit", "NAP",
            "Date Facture", "Date Paiement", "Pourcentage Paye", "Statut" )
        VALUES(NULL,?,?,?,?,?,?,?)""", bill_data)
    connection.commit()

    # "confirmed" table linking bills to parent jobs and clients
    cursor.execute("""
        CREATE TABLE confirmed (
        key INTEGER PRIMARY KEY,
        client_id INTEGER,
        job_id INTEGER,
        bill_id INTEGER,
        FOREIGN KEY(client_id) REFERENCES client(key),
        FOREIGN KEY(job_id) REFERENCES job(key),
        FOREIGN KEY(bill_id) REFERENCES bill(key)
        );""")

    confirmed_data = [(1, 1, 1),
                      (1, 2, 0),
                      (2, 3, 2),
                      (2, 4, 3),
                      (2, 5, 0),
                      (3, 6, 4),
                      (3, 7, 5),
                      ]

    cursor.executemany("""
        INSERT INTO confirmed ("key","client_id", "job_id", "bill_id")
        VALUES(NULL,?,?,?)""", confirmed_data)
    connection.commit()

    # "invoice_items table linking bills to parent jobs and clients
    cursor.execute("""
        CREATE TABLE invoice_items (
        key INTEGER PRIMARY KEY,
        Category VARCHAR,
        Start_Date VARCHAR,
        End_Date VARCHAR,
        Description VARCHAR,
        Quantity VARCHAR,
        Unit_Price VARCHAR,
        Subtotal VARCHAR,
        Invoice_id INTEGER,
        FOREIGN KEY (Invoice_id) REFERENCES bill (key)
        );""")

    invoice_data = [("Couverture Evenement", "15/10/2016", "15/10/2016",
                     "Suivi Evenement Bingo", "1", "500000", "500000", 1),

                    ("Shooting Photo", "15/10/2016", "15/10/2016",
                     "Photoshoot model", "1", "500000", "500000", 1),

                    ("Shooting Video", "15/10/2016", "15/10/2016",
                     "Video dans le Holi DZ", "1", "0", "0", 1),

                    ("3D Conception (Stand)", "15/10/2016", "15/10/2016",
                     "STAND ATPA", "1", "1500000", "1500000", 1),

                    ("Shooting Photo", "15/10/2016", "15/10/2016",
                     "Photoshoot model", "1", "500000", "500000", 2),

                    ("Shooting Video", "15/10/2016", "15/10/2016",
                     "Video dans le Holi DZ", "1", "0", "0", 2),

                    ("3D Conception (Stand)", "15/10/2016", "15/10/2016",
                     "STAND ATPA", "1", "1500000", "1500000", 2),

                    ("Shooting Photo", "15/10/2016", "15/10/2016",
                     "Photoshoot model", "1", "500000", "500000", 3),

                    ("Shooting Video", "15/10/2016", "15/10/2016",
                     "Video dans le Holi DZ", "1", "0", "0", 3),

                    ("3D Conception (Stand)", "15/10/2016", "15/10/2016",
                     "STAND ATPA", "1", "1500000", "1500000", 4)]

    cursor.executemany("""
        INSERT INTO invoice_items ("key","Category", "Start_Date", "End_Date",
                               "Description", "Quantity", "Unit_Price",
                               "Subtotal", "Invoice_id")
        VALUES(NULL,?,?,?,?,?,?,?,?)""", invoice_data)
    connection.commit()

    # Modification table that stores modification items to be linked with job
    cursor.execute("""
        CREATE TABLE debrief (
        key INTEGER PRIMARY KEY AUTOINCREMENT,
        Description VARCHAR,
        [Date_Created] TIMESTAMP,
        [Date_Done] TIMESTAMP,
        Parent TEXT,
        FOREIGN KEY (Parent) REFERENCES job ("Code Job")
        );""")
    connection.commit()
    debrief_data = [
                    ("Rajouter une chaise", "15/10/2016", "15/10/2016",
                     "J/002/003"),
                    ("Rajouter des tables", "15/10/2016", "",
                     "J/002/003"),
                    ("Rajouter du rooooose", "15/10/2016", "15/10/2016",
                     "J/002/004"),
                    ("Rajouter des poupenszzzz", "15/10/2016", "",
                     "J/002/003"),
                     ]

    cursor.executemany("""
        INSERT INTO debrief ("key","Description", [Date_Created], [Date_Done],
                            "Parent")
        VALUES(NULL,?,?,?,?)""", debrief_data)
    connection.commit()

    connection.close()


def del_selected(tabview, model, table=None):
    """Func to diplay QtGui.QMessageBox before deleting entry.

    Args:
        table (QtGui.QTableview): The first parameter.
        param2 (QtSql.QSqlTableModel): The second parameter.

    Returns:
        bool: The return value. True for success, False otherwise.
    """
    msg = QtGui.QMessageBox()
    msg.setIcon(QtGui.QMessageBox.Warning)
    msg.setText("Are you sure you want to delete this data ?")
    # msg.setInformativeText("To edit a client you must select a row.")
    msg.setWindowTitle("Delete")
    msg.setStandardButtons(QtGui.QMessageBox.Ok | QtGui.QMessageBox.Cancel)
    set_style(msg)
    result = msg.exec_()

    if result == QtGui.QMessageBox.Ok:
        if table == "job":
            try:
                # remove orders taoble sql and update
                L = tabview.selectionModel().selectedRows()
                for index in sorted(L):
                    print('Row %d is selected' % index.row())

                index = tabview.selectedIndexes()[1]
                id_us = str(tabview.model().data(index).toString())
                print "code : ", id_us
                cid = str(id_us[4])
                jid = str(id_us[8])

                print "cid: ", cid
                print "jid: ", jid
                print 'Yes.'
                # Delete link from orders table
                connection = sqlite3.connect(DB)
                cursor = connection.cursor()
                # id_us =  "\"" + id_us + "\""
                # cursor.execute("DELETE FROM orders WHERE Code Job = '%s';"
                #                % (id_us.strip()))

                cursor.execute("""
                        DELETE FROM orders
                        WHERE "client_id" = '%s' AND "job_id" = '%s' """ % (cid, jid))
                connection.commit()
                # UNCOMMENT BELOW TO DEBUG
                # cursor.execute("""SELECT * FROM orders""")
                # result_orders = cursor.fetchall()
                connection.close()
            except (IndexError, ValueError):

                sel_msg = QtGui.QMessageBox()
                sel_msg.setIcon(QtGui.QMessageBox.Information)

                sel_msg.setText("Oops! Selection Error:")
                sel_msg.setInformativeText("To delete a client you must select a row.")
                sel_msg.setWindowTitle("Input Error")

                sel_msg.setStandardButtons(QtGui.QMessageBox.Ok |
                                       QtGui.QMessageBox.Cancel)
                set_style(sel_msg)
                sel_msg.exec_()
                return



        # # UNCOMMENT SECTION FOR DEBUGGING /START/
        # print "result orders after deletion:"
        # for r in result_orders:
        #     print r

        indices = tabview.selectionModel().selectedRows()
        for index in sorted(indices):
            model.removeRow(index.row())
        lambda: model.removeRow(
            tabview.currentIndex().row())

    else:
        print 'No.'
    # warn.exec_()
    # w.close()


def update_combo_from_db(combo, column, table, displayfor=None):
    """Function update contents of ComboBox.

    Rewritten to take following arguments:
    `ComboBox`, `ColumnName`, `Table` in order to fill ComboBox
    with rows from `ColumnName` in `Table`.


    """
    # UNCOMMENT SECTION FOR DEBUGGING /START/
    # print "\n"
    # print "-" * 80
    # print "START OF FUNCTION update combo client"
    # UNCOMMENT SECTION FOR DEBUGGING /END/
    # if displayfor is not None:
    #     print "displayfor mode active>>", displayfor
    connection = sqlite3.connect(DB)
    cursor = connection.cursor()
    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    # cursor.execute("""SELECT * FROM client""")
    # result_all = cursor.fetchall()
    # print "result all:", result_all
    # # UNCOMMENT SECTION FOR DEBUGGING /END/

    # Unused part of the code, should try to reimplement it
    # cursor.execute("""
    #     SELECT DISTINCT "{col}" FROM "{tab}" """.format(col=column, tab=table))
    # result_key = cursor.fetchall()

    # cursor.execute("""SELECT DISTINCT "Organisation" FROM client""")
    # result_company = cursor.fetchall()

    cursor.execute("""SELECT * FROM client""")
    result_everything = cursor.fetchall()

    cursor.execute("""SELECT * FROM job""")
    result_everything_job = cursor.fetchall()

    if displayfor != "<Client>" and displayfor is not None:
        # print "Print keys of jobs linked to client"
        cursor.execute("""SELECT * FROM orders WHERE client_id = "{k}" """
                .format(k=displayfor[:3]))
        result_displayfor = cursor.fetchall()
        # print result_displayfor
        combo.clear()
        combo.addItem("<Projet a debriefer>")


    connection.close()

    # UNCOMMENT SECTION FOR DEBUGGING /START/
    # print "result_key:", result_key
    # print "result_company:", result_company
    # print "result_everything:", result_everything
    if table == "client":
        result = result_everything
    elif table == "job":
        result = result_everything_job
    else:
        raise NotImplementedError
    id_and_comp = []
    for x in result:
        # print "x>>> ", x
        seq = (x[1], x[3])
        combined = ' '.join(seq)
        id_and_comp.append(combined)
        # print "combined>>>", combined
    # print "id + company:", id_and_comp #[x[1] for x in result_everything],
    # UNCOMMENT SECTION FOR DEBUGGING /END/
    # print id_and_comp
    for p in id_and_comp:
        # print "p>>>", p
        if p is not None:
            # # UNCOMMENT SECTION FOR DEBUGGING /START/
            # print "unecoded", p
            # # UNCOMMENT SECTION FOR DEBUGGING /END/

            p = ucd.normalize('NFKD', p).encode('ascii', 'ignore')
            # UNCOMMENT SECTION FOR DEBUGGING /START/
            # print "encoded", p
            # UNCOMMENT SECTION FOR DEBUGGING /END/
        else:
            # # UNCOMMENT SECTION FOR DEBUGGING /START/
            # print p, "DID NOT ENCODE"
            # # UNCOMMENT SECTION FOR DEBUGGING /END/
            return
        for i in range(combo.count()):

            # # retrieve combox current items
            allitems = [str(combo.itemText(u))
                        for u in range(combo.count())]
            # UNCOMMENT SECTION FOR DEBUGGING /START/
            # print "All items:", allitems
            # UNCOMMENT SECTION FOR DEBUGGING /END/

            if p not in allitems:
                # # UNCOMMENT SECTION FOR DEBUGGING /START/
                # print "DID NOT ADD ITEM TO COMBO BOX"
                # # add the name of the client by list compreh. search
                # tup = [item for item in result_all if p in item]
                # print "code:", tup
                # # UNCOMMENT SECTION FOR DEBUGGING /END/

                # When in `displayfor` mode add only relevant jobs to combo, else
                # add them all
                if displayfor != "<Client>" and displayfor is not None:

                    # print "Should only display jobs for client >> %r" % displayfor[:3]
                    # print "p is >>", p

                    p_split = p[2:5]
                    p_split = p_split.strip()
                    # print "p split is >> %r" % p_split
                    # print "job matches client", p_split == displayfor[:3]
                    if p_split == displayfor[:3]:
                        # print "MATCH Added item to Combo modif!"
                        combo.addItem(p)

                else:
                    combo.addItem(p)
                # # UNCOMMENT SECTION FOR DEBUGGING /START/
                # print "ADDED ITEM %s TO COMBO BOX" % p
                # print "All items:", allitems
                # # UNCOMMENT SECTION FOR DEBUGGING /END/
    # print "All items:", allitems

    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    # print "END OF FUNCTION update combo client"
    # print "-" * 80
    # print "\n"
    # # UNCOMMENT SECTION FOR DEBUGGING /END/


def input_window(tab=None, model=None, combo=None, mode=None,
    fields=None, tabview=None, cols_to_fetch=None):
    """Create modal dialog w/ len(fields) rows of QPushButton & QLineEdit pairs.

    Args:
        tab     (SQLITE table):     -Table to read/write.
        model (QSqlTableModel):     -Tableview model to update.
        combo      (QComboBox):     -combo box data to be fetched.
        fields          (list):     -List of strings to generate buttons.
                                     Strings must match col names in SQL table
        mode             (str):     -New mode or edit mode for SQLITE scripts.
                                     If edit mode selected, we will display
                                     data to be modified in QLineEdit.
                                     If undefined or new, it will show empty...
    Returns:
        bool: The return value. True for success, False otherwise.
    """
    def get_text_input(var=None):
        text, ok = QtGui.QInputDialog.getText(
            var, 'Text Input Dialog', 'Enter Information:')
        if ok:
            var.setText(str(text))

    def cancel():
        win.close()

    def update_db_wrapper():
        col_dict = {}
        for key in fields:

            # print "lines[key].text():", lines[key].text()
            if key == "Code Client" or key == "Code Job":
                print "key is BAD: ", key
                continue
            elif str(lines[key].text()):
                print "key is GOOD: ", key

                col_dict[key] = str(lines[key].text())
            else:
                return

        # print "\n dict passed to update_db: "
        # print col_dict
        # print "row passed to update_db: ", key2read

        update_db(table=tab, row=key2read,
                  d=col_dict, close_win=win)
        model.select()

    def add_to_db_wrapper():
        #######################################################################
        # Part that updates DB entries
        #######################################################################
        col_dict = {}
        # print "Fields"
        # print fields

        for key in fields:
            # print "lines[key].text():", lines[key].text()
            if key == "Code Client" or key == "Code Job":
                # print "key is BAD: ", key
                continue
            elif str(lines[key].text()):
                # print "key is GOOD: ", key

                col_dict[key] = str(lines[key].text())
            else:
                return
        if tab == "client":
            # add stuff here if client
            col_dict["Factures Payees"] = "0/0"
            col_dict["Solde Compte"] = "DA 0.00"
            #col_dict["Factures Recues"] = 0
        # # UNCOMMENT SECTION FOR DEBUGGING /START/
        # print "dict passed to add_to_db_v2: "
        # print col_dict
        # # UNCOMMENT SECTION FOR DEBUGGING /END/
        add_to_db_v2(table=tab, d=col_dict, close_win=win)
        #######################################################################

        #######################################################################
        # Part that updates cc number live
        #######################################################################
        # UPDATE CLIENT CODE
        #####################
        connection = sqlite3.connect(DB)
        cursor = connection.cursor()
        cursor.execute("""SELECT * FROM client""")
        result_all_clients = cursor.fetchall()
        # # UNCOMMENT SECTION FOR DEBUGGING /START/
        # print "-" * 80
        # print "-" * 80
        # print("fetchall before cc update:")
        # for r in result_all_clients:
        #     print(r)
        # # UNCOMMENT SECTION FOR DEBUGGING /END/
        result_cc = [i[0] for i in result_all_clients]
        # print "result_cc = ", result_cc
        result_str = range(len(result_cc))
        int_str_tuple = result_str
        # print "result_str", result_str

        for idx, val in enumerate(result_cc):
            # # UNCOMMENT SECTION FOR DEBUGGING /START/
            # print "idx=", idx
            # # UNCOMMENT SECTION FOR DEBUGGING /END/
            count = "00000000" + str(val)  # +"\'"
            count = count[-3:]

            # count = "\"" + count + "\""
            # print "concatenated:", count
            result_str[idx] = count
            int_str_tuple[idx] = (count, val)
        # print "result_str", result_str
        if tab == "client":
            cursor.executemany("""
                UPDATE client SET "Code Client" = ?
                WHERE "key" = ?""", int_str_tuple)
            # end of update cc
            connection.commit()
            connection.close()
            ###################################################################

        elif tab == "job":
            ###################################################################
            # Create entry in table ORDERS to keep track of who ordered what.
            ###################################################################
            item = [str(combo.currentText())]
            # print
            # print "<<<item in combo before>>>", item

            item = item[:3]
            # print "<<<item in combo after>>>", item
            # print
            # fetch last job column
            connection = sqlite3.connect(DB)
            cursor = connection.cursor()
            cursor.execute("""SELECT DISTINCT "key" FROM job""")
            result_key = cursor.fetchall()
            connection.close()
            # print "All jobs in NewJobIptDia: ", result_key
            last = max(result_key)
            last = last[0]
            print "last: ", last
            print "should insert in row last+1: ", last + 1
            print "<<<item[0]>>>= ", item[0][:3]

            # set values in dict
            m = "<Select Client>"
            orders_dict = {"key": "NULL",
                           "client_id": item[0][:3],
                           "job_id": last,
                           }
            if item != m:
                # call method to update db#####################################
                add_to_db_v2(table="orders", d=orders_dict, close_win=win)
                ###############################################################

                ###############################################################
                # UPDATE JOB CODE AND ORDERS TABLE
                ###############################################################
                # UPDATE JOB CODE
                #################
                item[0] = "00000000" + str(item[0][:3])  # +"\'"
                item[0] = item[0][-3:]
                last = "00000000" + str(last)  # +"\'"
                last = last[-3:]
                code_job_str = "J/" + item[0] + "/" + last
                params = [code_job_str, last]
                print "params>>>: ", params

                connection = sqlite3.connect(DB)
                cursor = connection.cursor()
                cursor.execute("""
                    UPDATE job SET "Code Job" = ?
                    WHERE "key" = ?""", params)
                connection.commit()
                connection.close()

            tabview.model().select()

        else:
            raise NotImplementedError

            # cursor.execute("""SELECT * FROM job""")
            # result_all_jobs = cursor.fetchall()
            # for i in result_cc:
            #     # # UNCOMMENT SECTION FOR DEBUGGING /START/
            #     # print "i=", i
            #     # # UNCOMMENT SECTION FOR DEBUGGING /END/
            #     count = "00000000" + str(i)  # +"\'"
            #     count = count[-3:]
            #     count = "\"" + count + "\""
            #     # print "concatenated:", count
            #     cursor.execute("""
            #         UPDATE client SET "Code Client" = {s}
            #         WHERE "key" = {k}""" .format(s=count, k=i))
            #     # end of update cc
            # connection.commit()

        connection.close()
        model.select()

    # dict "existing_data" maps current data in db to qline edits for show
    existing_data = {}
    code = "---not yet created---"
    text = "Code "
    what = "Something went wrong...."
    window_info = "Information Input Window"

    for k in cols_to_fetch:
        existing_data[k] = ""
    if mode == "edit":
        # for i in sorted(indexes):
            # print('Row %d is selected' % i.row())
            # print "col"
        try:
            L = tabview.selectionModel().selectedRows()
            for index in sorted(L):
                print('Row %d is selected' % index.row())
            # uncomment for debug, view number of columns in table
            # (stored in all_indices)
            # all_indices = len(tabview.selectedIndexes()[:])
            # print "ALL INDICES: ", all_indices

            # cols_to_fetch = [1, 2, 3, 4, 7]
            # print cols_to_fetch
            for k in cols_to_fetch:
                print "k:", k
                i = tabview.selectedIndexes()[k]
                existing_data[k] = str(tabview.model().data(i).toString())
                if (k == 0):
                    key2read = int(existing_data[k])

                    # key2read = "\'" + str(key2read) + "\'"
                    # print "key2read: ", key2read
                # print "Existing data: ", existing_data[k]
            print existing_data

        except (IndexError, ValueError):

            msg = QtGui.QMessageBox()
            msg.setIcon(QtGui.QMessageBox.Information)

            msg.setText("Oops! Selection Error:")
            msg.setInformativeText("To edit a client you must select a row.")
            msg.setWindowTitle("Input Error")

            msg.setStandardButtons(QtGui.QMessageBox.Ok |
                                   QtGui.QMessageBox.Cancel)
            set_style(msg)
            msg.exec_()
            return
        # small section to format window according to table being acessed
        if (tab == 'client'):
            code = "00" + str(existing_data[0])
            what = ' Client'
            window_info = "Client Information Input Window"

        elif (tab == 'job') or (tab == 'bill'):
            code = str(existing_data[1])
            what = 'Job'
            window_info = "Job Information Input Window"
        else:
            code = "---not yet created---"
            return

    win = QtGui.QDialog()
    layout = QtGui.QGridLayout()
    layout.setVerticalSpacing(5)
    layout.setHorizontalSpacing(5)
    label = QtGui.QLabel()
    label.setText(
        "Please enter information by clicking on the buttons on the right.")
    label.setSizePolicy(QtGui.QSizePolicy.Expanding,
                        QtGui.QSizePolicy.Preferred)
    label.setFont(QtGui.QFont('', 12))
    label.setAlignment(QtCore.Qt.AlignCenter)
    layout.addWidget(label, 0, 0, 1, 2)

    if (tab == 'client'):
        cols_to_fetch = cols_to_fetch[2:]
        what = " Client"
        text += what
        # i is rwos to offset (number of rows before start putting rows
        # of buttons and qlineedits, currently == 2 because two labels
        i = 2

        window_info = "Information Input Window"
    elif (tab == 'job'):
        cols_to_fetch = cols_to_fetch[1:]
        what = " Job"
        text += what
        # i is rwos to offset (number of rows before start putting rows
        # of buttons and qlineedits, currently == 3 because two labels and
        # a widget come before button fields...
        #
        i = 3
        tree_selector = init_tree_view()
        layout.addWidget(tree_selector, 2, 0, 1, 2)

    else:
        print tab, mode
        raise NotImplementedError
    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    # for k in cols_to_fetch:
    #     existing_data[k] = ""

    # print "cols_to_fetch: ", cols_to_fetch
    # print "fields:", fields
    # # UNCOMMENT SECTION FOR DEBUGGING /END/
    input_map = dict(zip(fields, cols_to_fetch))
    # print "input map", input_map
    buttons = {}
    lines = {}
    # create buttons
    for key in fields:
        if key != ("Code Job" or "Code Client"):
            buttons[key] = QtGui.QPushButton(key)
            print existing_data[input_map[key]]
            db_data = existing_data[input_map[key]]
            lines[key] = QtGui.QLineEdit(db_data)
            lines[key].setReadOnly(True)
            buttons[key].clicked.connect(
                partial(get_text_input, var=lines[key]))
            # # UNCOMMENT SECTION FOR DEBUGGING /START/
            # print buttons[key]
            # print lines[key]
            # # UNCOMMENT SECTION FOR DEBUGGING /END/
            layout.addWidget(buttons[key], i, 0)
            layout.addWidget(lines[key], i, 1)
            i += 1
        else:
            continue

    label2 = QtGui.QLabel()
    label2.setText("{txt}: {cc}".format(txt=text, cc=code))
    label2.setFont(QtGui.QFont('', 12))
    label2.setAlignment(QtCore.Qt.AlignCenter)
    layout.addWidget(label2, 1, 0, 1, 2)

    btnDone = QtGui.QPushButton("Add" + what)
    btnCancel = QtGui.QPushButton("Cancel Operation")
    btnBox = QtGui.QHBoxLayout()
    btnBox.addWidget(btnDone)
    btnBox.addWidget(btnCancel)
    layout.addItem(btnBox, i + 1, 1)

    btnCancel.clicked.connect(cancel)
    if mode == "edit":
        btnDone.clicked.connect(update_db_wrapper)
    else:
        btnDone.clicked.connect(add_to_db_wrapper)

    win.setLayout(layout)
    win.setWindowTitle(window_info)
    set_style(win)
    if (tab == 'client'):
        height = 80 + 50 * len(fields)
    elif (tab == 'job'):
        height = 245 + 50 * len(fields)
    else:
        raise NotImplementedError
    # win.setFixedHeight(height)
    # win.setFixedWidth(400)
    win.exec_()


def get_jobs_tabview(combo, tab_view, model2, showall=None):
    """Fetch data from ComboBox & display corresponding rows in QTableView.

    Thia function will show only jobs of selected client by hiding all other
    rows than the ones linked by `orders` table.
    """
    # get items from client db
    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    # print "\n"
    # print "-" * 80
    # print "START METHOD 'get_jobs_tabview'... \n"
    # # UNCOMMENT SECTION FOR DEBUGGING /END/
    connection = sqlite3.connect(DB)
    cursor = connection.cursor()
    cursor.execute("""SELECT * FROM orders""")
    result_orders = cursor.fetchall()
    connection.close()
    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    # print "result orders:"
    # for r in result_orders:
    #     print r
    # # UNCOMMENT SECTION FOR DEBUGGING /END/
    # cursor.execute("""SELECT DISTINCT "Code Client" FROM client""")
    # result_key = cursor.fetchall()
    # for r in result_key:
    #     print "result:", r
    if showall is None:
        selected_item = [str(combo.currentText())]
        selected_item = selected_item[:3]
    elif showall is True:
        selected_item = "<View All>"
    elif showall is False:
        selected_item = "<Select Client>"

    # print "selected_item[0]= ", selected_item[0]
    m = "<Select Client>"
    n = "<View All>"

    # print "Condition input != <Select Clients> evaluates:", selected_item[0]
    # != m
    # print "selected_item>>>", selected_item
    # print "selected_item[0][:3]>>>", selected_item[0][:3]
    # code here to take p and fetch all row ids from
    # orders table to display relevant rows from jobs table
    # load db into tableJobs
    tab_view.setModel(model2)
    tab_view.setSortingEnabled(False)
    tab_view.setEditTriggers(QtGui.QAbstractItemView.NoEditTriggers)
    # trick to auto resize columns
    tab_view.setVisible(False)
    tab_view.resizeColumnsToContents()
    tab_view.resizeRowsToContents()
    hh = tab_view.horizontalHeader()
    hh.setStretchLastSection(True)
    tab_view.setVisible(True)
    tab_view.setShowGrid(False)
    # tab_view.setSortingEnabled(True)
    tab_view.model().layoutChanged.emit()
    count = model2.rowCount()
    # print "count: ", count
    # print "client selected:", selected_item[0]
    model2.select()
    if selected_item[0] == m:
        for i in range(count):
            tab_view.hideRow(i)
            # print "entered where selected item= select client"
            # tab_view.setColumnWidth ( int column, int width)

    elif selected_item[0] == n:
        for i in range(count):
            tab_view.showRow(i)
    else:
        print "selected_item[0][:3]>>", selected_item[0][:3]

        # Explained in words: For each i, v in a enumerated list of L
        # (that makes i the element's position in the enumerated list
        # and v the original tuple) check if the tuple's first element
        # is 53, if so, append the result of the code before'for' to a
        # newly created list, here: i. It could also be my_function(i, v)
        # or yet another list comprehension. Since your list of tuples only has
        # one tuple with 53 as first value, you will get a list with one elemt.
        if selected_item[0][:3] != "<":
            tup2 = [i for i, v in enumerate(result_orders) if v[
                1] == int(selected_item[0][:3])]
            # print "list of jobs keys linked to client :", tup2
            for i in range(count):
                tab_view.hideRow(i)
            for i in tup2:
                tab_view.showRow(i)
            tab_view.model().layoutChanged.emit()
            # # UNCOMMENT SECTION FOR DEBUGGING /START/
            # print "Client<", selected_item[0], ">selected"
            # print "END OF METHOD 'get_jobs_tabview'... \n"
            # print "-" * 80
            # print "\n"
            # # UNCOMMENT SECTION FOR DEBUGGING /END/


def get_bills(combo, tab_view, model3, table, showall=None):
    """Fetch data from ComboBox & display corresponding rows in QTableView.

    Thia function will show only bills of selected client or selected job
    by hiding all other rows than the ones linked by `confirmed` table.
    """
    # get items from client db
    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    # print "\n"
    # print "-" * 80
    # print "START METHOD 'get_jobs_tabview'... \n"
    # # UNCOMMENT SECTION FOR DEBUGGING /END/
    connection = sqlite3.connect(DB)
    cursor = connection.cursor()
    cursor.execute("""SELECT * FROM "{tab}" """.format(tab=table))
    result_orders = cursor.fetchall()
    connection.close()
    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    print "result %s:" % table
    for r in result_orders:
        print r
    # # UNCOMMENT SECTION FOR DEBUGGING /END/
    # cursor.execute("""SELECT DISTINCT "Code Client" FROM client""")
    # result_key = cursor.fetchall()
    # for r in result_key:
    #     print "result:", r
    if showall is True:
        selected_item = "<View All>"
    else:
        selected_item = [str(combo.currentText())]

    # elif showall is False:
    #     selected_item = "<Select Client>"

    print "selected_item[0]= ", selected_item[0]

    m = "<Select Client>"
    n = "<View All>"
    o = "<Select Job>"

    # print "Condition input != <Select Clients> evaluates:", selected_item[0]
    # != m

    # code here to take p and fetch all row ids from
    # orders table to display relevant rows from jobs table
    # load db into tableJobs
    tab_view.setModel(model3)
    tab_view.setSortingEnabled(False)
    tab_view.setEditTriggers(QtGui.QAbstractItemView.NoEditTriggers)
    # trick to auto resize columns
    tab_view.setVisible(False)
    tab_view.resizeColumnsToContents()
    tab_view.resizeRowsToContents()
    hh = tab_view.horizontalHeader()
    hh.setStretchLastSection(True)
    tab_view.setVisible(True)
    tab_view.setShowGrid(False)
    # tab_view.setSortingEnabled(True)
    tab_view.model().layoutChanged.emit()
    count = model3.rowCount()
    # print "count: ", count
    # print "client selected:", selected_item[0]
    model3.select()
    if selected_item[0] == (m or o):
        for i in range(count):
            tab_view.hideRow(i)
            # tab_view.setColumnWidth ( int column, int width)

    elif selected_item[0] == n:
        for i in range(count):
            tab_view.showRow(i)
    else:
        # Explained in words: For each i, v in a enumerated list of L
        # (that makes i the element's position in the enumerated list
        # and v the original tuple) check if the tuple's first element
        # is 53, if so, append the result of the code before'for' to a
        # newly created list, here: i. It could also be my_function(i, v)
        # or yet another list comprehension. Since your list of tuples only has
        # one tuple with 53 as first value, you will get a list with one elemt.
        tup2 = [i for i, v in enumerate(result_orders) if v[
            1] == int(selected_item[0])]
        print "list of jobs keys linked to client :", tup2
        for i in range(count):
            tab_view.hideRow(i)
        for i in tup2:
            tab_view.showRow(i)
        tab_view.model().layoutChanged.emit()
        # # UNCOMMENT SECTION FOR DEBUGGING /START/
        # print "Client<", selected_item[0], ">selected"
        # print "END OF METHOD 'get_jobs_tabview'... \n"
        # print "-" * 80
        # print "\n"
        # # UNCOMMENT SECTION FOR DEBUGGING /END/


def new_bill_input_dialog(combo=None, table=None, model=None):
    """Launch billQdialog.py and retrieve billing info.

    This function will show a dialog for the user to input invoice data
    into a GUI and be able to store it in a database.
    """
    def cancel():
        win.close()

    def update_db_wrapper():

        # cursor.execute("""
        #     CREATE TABLE confirmed (
        #     key INTEGER PRIMARY KEY,
        #     client_id INTEGER,
        #     job_id INTEGER,
        #     bill_id INTEGER,
        #     FOREIGN KEY(client_id) REFERENCES client(key),
        #     FOREIGN KEY(job_id) REFERENCES job(key)
        #     FOREIGN KEY(bill_id) REFERENCES bill(key)
        #     );""")

        bill_dict = {
            "key": "NULL",
            "Produit": "Shooting",
            "Brief": "Pub Pepsi",
            'Date Debut': "15/11/2016",
            'Date Fin': "15/12/2016",
            "Prix": "2500000",
            "Somme Payee": "1250000",
        }

        add_to_db_v2(table="bill", d=bill_dict, close_win=win)

        # create entry in link table ORDERS to keep track of who ordered what
        item = [str(combo.currentText())]

        # fetch last job column
        connection = sqlite3.connect(DB)
        cursor = connection.cursor()
        cursor.execute("""SELECT DISTINCT "key" FROM bill""")
        result_key = cursor.fetchall()
        connection.close()
        print "All jobs in NewBillIptDia: ", result_key
        last = max(result_key)
        last = last[0]
        print "last: ", last
        print "should insert in row last+1: ", last + 1
        print "item[0]= ", item[0]

        # set values in dict
        m = "<Select Client>"
        bill_dict = {"key": "NULL",
                     "client_id": item[0],
                     "job_id": last,
                     }
        if item != m:
            # call method to update db
            add_to_db_v2(table="confirmed", d=bill_dict, close_win=win)
        table.model().select()

    def txtchanged(text, k=None):
        # print "%s: " % (k.objectName())
        in_fields_2_data[k] = text
        in_fields_2_data_str[str(k.objectName())] = text

        print "raw %r:" % str(k.objectName())

        in_fields_2_pLbls[k].setText(in_fields_2_data_str[str(k.objectName())])

        # if k.objectName() == u'leClientName':
        #     win.labelClient.setText(in_fields_2_data_str[str(k.objectName())])
        #     print "!!!entered!!!"

        # print "k.objectName() is type", type(k.objectName())
        # print "saved data: ", in_fields_2_data[k]

    def datechanged(date, k=None):
        in_fields_2_data[k] = date
        in_fields_2_data_str[str(k.objectName())] = str(date.toPyDate())
        print "raw %r:" % str(k.objectName())
        print "k.objectName() is type", type(k.objectName())
        in_fields_2_pLbls[k].setText(str(date.toPyDate()))

    def numchanged(num, k=None):
        in_fields_2_data[k] = num
        in_fields_2_data_str[str(k.objectName())] = num
        in_fields_2_pLbls[k].setText(in_fields_2_data_str[str(k.objectName())])

    def cb_selectchanged(index, k=None):
        print "index of cb: ", index, "value: ", k.currentText()
        in_fields_2_data[k] = k.currentText()
        in_fields_2_data_str[str(k.objectName())] = str(in_fields_2_data[k])
        in_fields_2_pLbls[str(k)].setText(in_fields_2_data_str[str(k.objectName())])

    def add_2_inv_items():
        pass

    def preview_inv_tex():
        print
        print "Invoice Preview:"
        print "=" * 80
        print
        print "ObjectName dict{}", len(in_fields_2_data_str), "item(s) long."
        for key in in_fields_2_data_str:
            print key, ":", in_fields_2_data_str[key]
        print "END OF Invoice Preview:"
        print "=" * 80

        generate_latex(inv_data=in_fields_2_data_str, doc="invoice")
        import subprocess
        # first = "C:/Users/Nuts/Desktop/Py/dev/kiwi/gui/"
        # second = ".texbuild/jinja2build/jinja2-testOut.pdf"
        # pdf = first + second
        rel_pdf = ".texbuild/jinja2build/jinja2-testOut.pdf"
        pdf = os.path.join(os.sep, os.path.dirname(__file__), rel_pdf)
        acrobatPath = r'SumatraPDF/SumatraPDF.exe'
        subprocess.Popen("%s %s" % (acrobatPath, pdf))

    def preview_invoice_docx():
        print
        print "Invoice Preview (passed args):"
        print "=" * 80
        print
        print "ObjectName dict{}", len(in_fields_2_data_str), "item(s) long."
        for key in in_fields_2_data_str:
            print key, ":", in_fields_2_data_str[key]
        print "END OF Invoice Preview (passed args):"
        print "=" * 80

        generate_docx(inv_data=in_fields_2_data_str, doc="invoice")

    def tex_docx_btnstate(b):
        if b.text() == "LaTeX":
            if b.isChecked():
                print b.text() + " is selected"
                win.btnGenPreview.clicked.connect(preview_inv_tex)
            else:
                print b.text() + " is deselected"
                win.btnGenPreview.clicked.disconnect()

        if b.text() == "MS Word":
            if b.isChecked():
                print b.text() + " is selected"
                win.btnGenPreview.clicked.connect(preview_invoice_docx)
            else:
                print b.text() + " is deselected"
                win.btnGenPreview.clicked.disconnect()

    def inner_print_test():
        print
        print "Locked/Unlocked"
        print

    try:
        _fromUtf8 = QtCore.QString.fromUtf8
    except AttributeError:
        def _fromUtf8(s):
            return s
    try:
        _encoding = QtGui.QApplication.UnicodeUTF8

        def _translate(context, text, disambig):
            return QtGui.QApplication.translate(
                context, text, disambig, _encoding)
    except AttributeError:
        def _translate(context, text, disambig):
            return QtGui.QApplication.translate(
                context, text, disambig)

    win = Ui_Dialog()
    dialog = QtGui.QDialog()
    win.setupUi(dialog)
    set_style(dialog)

    print "\n<<Successfully Added Model to Tree selector>>\n"
    tree_selector = init_tree_view()
    win.horizontalLayout_18.addWidget(tree_selector)

    print "\n Bill dialog opened."
    # load model into view for invoice items
    win.invoiceItemsTable.setModel(model)  # load db into tableClients
    win.invoiceItemsTable.setEditTriggers(
        QtGui.QAbstractItemView.NoEditTriggers)
    # trick to auto resize columns, hide, configure then show.
    win.invoiceItemsTable.setVisible(False)
    win.invoiceItemsTable.resizeColumnsToContents()
    win.invoiceItemsTable.resizeRowsToContents()
    hh = win.invoiceItemsTable.horizontalHeader()
    hh.setStretchLastSection(True)
    win.invoiceItemsTable.setShowGrid(False)
    win.invoiceItemsTable.setSortingEnabled(False)
    win.invoiceItemsTable.setVisible(True)
    count = model.rowCount()
    for i in range(count):
        win.invoiceItemsTable.hideRow(i)

    # define fields that user inputs in dict to loop over when the text changes
    input_keys = [win.leNumBill, win.leCity, win.leClientName,
                  win.leProduct, win.deCreationDate, win.cbItemCat,
                  win.deItemDateStart, win.deItemDateEnd, win.sbItemQty,
                  win.leItemUnitPrice, win.leItemDescription]
    # input_keys   = [win.leNumBill, win.leCity, win.leClientName]
    # implement dict mapping of line below
    # preview_labels = [win.labelClient, win.labelCodeFac, win.labelProduit,
    #                   win.labelDate]
    # input_keys   = [win.leNumBill, win.leCity, win.leClientName]

    in_fields_2_data = {key: '' for key in input_keys}
    in_fields_2_data_str = {str(key.objectName()): '<>' for key in input_keys}
    in_fields_2_pLbls = {win.leClientName: win.labelClient,
                         win.leNumBill: win.labelCodeFac,
                         win.leProduct: win.labelProduit,
                         win.deCreationDate: win.labelDate}
    # # UNCOMMENT SECTION FOR DEBUGGING /START/
    # print "-" * 60
    # print "Before Ok:"
    # print "Object dict{}", len(in_fields_2_data), "item(s) long: "
    # for key in in_fields_2_data:
    #     print key, ":", in_fields_2_data[key]
    # print
    # print "ObjectName dict{}", len(in_fields_2_data_str), "item(s) long: "
    # for key in in_fields_2_data_str:
    #     print key, ":", in_fields_2_data_str[key]
    # print "-" * 60
    # # UNCOMMENT SECTION FOR DEBUGGING /END/

    for key in in_fields_2_data:
        # Loop over elements in dict and according to their class
        # (QLineEdit, QDateEdite, QSpinBox etc.) extract data and
        # store in corresponding value of dict

        # # UNCOMMENT SECTION FOR DEBUGGING /START/
        # print "key : %s" % key
        # print "key.objectName(): ", str(key.objectName())
        # print "value: ", in_fields_2_data[key]
        # # UNCOMMENT SECTION FOR DEBUGGING /END/

        if str(key.objectName()).startswith('le'):
            key.textChanged.connect(partial(txtchanged, k=key))

        elif str(key.objectName()).startswith('de'):
            key.dateChanged.connect(partial(datechanged, k=key))

        elif str(key.objectName()).startswith('sb'):
            key.valueChanged.connect(partial(numchanged, k=key))

        elif str(key.objectName()).startswith('cb'):
            key.currentIndexChanged.connect(partial(cb_selectchanged, k=key))

        else:
            print "not valid format"
    # Business code
    win.rbLatex.toggled.connect(partial(tex_docx_btnstate, b=win.rbLatex))
    win.rbWord.toggled.connect(partial(tex_docx_btnstate, b=win.rbWord))

    result = dialog.exec_()

    # mapping variables from dict to labels as user types to show preview...
    win.labelClient.setText(in_fields_2_data_str[u'leClientName'])
    win.labelCodeFac.setText(in_fields_2_data_str["leNumBill"])
    win.labelProduit.setText(in_fields_2_data_str["leProduct"])
    win.labelDate.setText(str(in_fields_2_data_str["deCreationDate"]))

    print "-" * 80
    print "Bill dialog closed with a",

    if result == 1:
        print "`Ok`"
        # do stuff here like wirte a docx. document using that gathered data
    else:
        print "`Cancel`"

    # UNCOMMENT SECTION FOR DEBUGGING /START/
    # print "-" * 60
    # print "After Ok:"
    # print "Pressed `Ok`"
    # print "Object dict{}", len(in_fields_2_data), "item(s) long."
    # for key in in_fields_2_data:
    #     print key, ":", in_fields_2_data[key]

    # print

    # print "ObjectName dict{}", len(in_fields_2_data_str), "item(s) long."
    # for key in in_fields_2_data_str:
    #     print key, ":", in_fields_2_data_str[key]
    # print "-" * 60
    # UNCOMMENT SECTION FOR DEBUGGING /START/

    # btnDone.clicked.connect(update_db_wrapper)
    # btnCancel.clicked.connect(cancel)


def generate_docx(inv_data=None, doc=None):
    """GEN .docx document based on args and a template, and python-docx.

    This function will programatically generate word .docx documents replacing
    passed Args to display invoices and estimates based on templates.
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT

    print "=" * 80
    print "inv_data:"
    for key in inv_data:
        print key, ":", inv_data[key]
    print "=" * 80

    items = {}
    print "items{}:"
    for key in inv_data.keys():
        items[key] = str(inv_data[key])
        print "%r: %r" % (key, items[str(key)])

    print type(items['leNumBill'])
    document = Document()
    p = document.add_paragraph('Facture N: ')
    p.add_run(items['leNumBill']).bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph('Client: ')
    p.add_run(items['leClientName']).bold = True

    # paragraph_format = p.paragraph_format
    # tab_stops = paragraph_format.tab_stops
    # tab_stop = tab_stops.add_tab_stop(Pt(100), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run('\t\t\t\t\t\t\t Alger le: ')
    p.add_run(items['deCreationDate']).bold = True

    p = document.add_paragraph('Produit: ')
    p.add_run(items['leProduct']).bold = True

    table = document.add_table(rows=1, cols=4)
    table.style = 'TableGrid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Designation'
    hdr_cells[1].text = 'Quantite'
    hdr_cells[2].text = 'Prix'
    hdr_cells[3].text = 'Montant'
    document.save('demo.docx')


def generate_latex(inv_data=None, doc=None, inv_items=None):
    """GEN Latex document based on args and a template, and jinja2.

    This function will programatically generate latex .tex documents replacing
    passed Args to display invoices and estimates based on templates.
    """
    # add arguments to pass...
    import jinja2
    from jinja2 import Template
    import io

    print "=" * 80
    print "inv_data:"
    for key in inv_data:
        print key, ":", inv_data[key]
    print "=" * 80

    # function used to split number every 3rd char
    # and add a comma
    def splitAt(w, n):
        for i in range(0, len(w), n):
            yield w[i: i + n]

    latex_jinja_env = jinja2.Environment(
        block_start_string='\BLOCK{',
        block_end_string='}',
        variable_start_string='\VAR{',
        variable_end_string='}',
        comment_start_string='\#{',
        comment_end_string='}',
        line_statement_prefix='%%',
        line_comment_prefix='%#',
        trim_blocks=True,
        autoescape=False,
        loader=jinja2.FileSystemLoader(os.path.abspath('.'))
    )

    if doc == "invoice":
        # Define variables...
        project = "./"

        build_d = "{}.texbuild/jinja2build/".format(project)
        out_fil = "{}jinja2-testOut".format(build_d)
        template = latex_jinja_env.get_template('kiwi-invoice-template-1.tex')

        # # user inputs hardcoded for testing (to delete afterwards...)
        # numfac = 'FAC000-000'
        # client_name = 'ALSTOM ALGERIE'
        # city = 'Alger'
        # date_bill_creation = '14/07/2016'
        # product = 'Reportage Photo Constantine'
        # nap = "Cent Quatre Vingt Sept Mille Deux Cent"
        # category = "Couverture Evenement"
        # qty = "1"
        unit_price = "160000"

        # # user inputs using variables
        numfac = inv_data["leNumBill"]
        client_name = inv_data["leClientName"]
        city = inv_data["leCity"]
        date_bill_creation = inv_data["deCreationDate"]
        product = inv_data["leProduct"]
        nap = "Cent Quatre Vingt Sept Mille Deux Cent"
        category = "Couverture Evenement"
        qty = "1"
        unit_price = "160000"

        inv_items = {
            "Category": "Couverture Evenement",
            "Start_Date": "15/10/2016",
            "End_Date": "15/10/2016",
            "Description": "Suivi Evenement Bingo",
            "Quantity": "1",
            "Unit_Price": "500000",
            "Subtotal": "500000",
            "Invoice_id": "1"}

        total_price = str(int(qty) * int(unit_price))
        numbers = [unit_price, total_price]

        for idx, val in enumerate(numbers):
            numbers[idx] = ",".join(splitAt(val, 3))

        # create the build directory if not exisiting
        if not os.path.exists(build_d):
            os.makedirs(build_d)

        # print(template.render(section1='Long Form', section2='Short Form'))
        tex_code = template.render(numfac=numfac,
                                   client_name=client_name,
                                   city=city,
                                   date_bill_creation=date_bill_creation,
                                   nap=nap,
                                   category=category,
                                   qty=qty,
                                   unit_price=numbers[0],
                                   total_price=numbers[1],
                                   product=product,
                                   inv_items=inv_items)
        # saves tex_code to outpout file
        with io.open(out_fil + ".tex", "w", encoding='utf8') as f:
            f.write(tex_code)

        os.system("pdflatex -output-directory {} {}".format(
            os.path.realpath(build_d), os.path.realpath(out_fil)))
        # shutil.copy2(out_fil+".pdf", os.path.dirname(
        #                                os.path.realpath(tex_code)))
    elif doc == "estimate":
        raise NotImplementedError
    elif (doc is not None) and ((doc != "invoice") or (doc != "estimate")):
        print "Arg `%s` is invalid, please use `invoice or `estimate`." % doc
    else:
        print "Error in generate_latex(inv_data, doc, inv_items)"


def init_tree_view(mode=None):
    """Function to initialise the treeView found in home window."""
    # if table in db does not exist, pop up window showing treeview
    # to allow user an initia description of their products and Services
    # as a hierarchical tree. Store thin in db
    connection = sqlite3.connect(DB)
    cursor = connection.cursor()
    # DB has three main tables and 2 linking tables.
    # Code copied and pasted 3 times... need to find sol, in te mean time it
    # works...
    cursor.execute("""
    SELECT count(*) FROM sqlite_master
    WHERE type='table' AND name='product_structure';""")
    result = cursor.fetchall()
    connection.close()
    result = result[0][0]
    # print result
    if result is False:
        print "will create table here..."
        # product structure tablee ...
        # cursor.execute("""
        #     CREATE TABLE product_structure (
        #     key INTEGER PRIMARY KEY AUTOINCREMENT,
        #     name VARCHAR(30),
        #     flat_price REAL,
        #     rate_price REAL,
        #     parent_key INTEGER NULL REFERENCES product_structure (key)
        #     ;""")
    # # product structure tablee ...
    # cursor.execute("""
    #     CREATE TABLE client (
    #     "key" INTEGER PRIMARY KEY AUTOINCREMENT,
    #     "Code Client" VARCHAR(30),
    #     "Nom" VARCHAR(30),
    #     "Organisation" VARCHAR(30),
    #     "Telephone" VARCHAR(30),
    #     "Factures Payees" INTEGER(1),
    #     "Factures Recues" INTEGER(1),
    #     "Email" VARCHAR(40),
    #     UNIQUE ("key") )

    reload(treeselector_)

    if mode == 'edit':
        tree_selector = treeselector_.Ui_()
        print 'created edit tree '
    else:
        tree_selector = treeselector_.Ui_layoutWidget()
        print 'created selection tree only'

    # need a function that maps sqlite table to  following lines
    rootNode   = treeselector_.Node("Products")
    childNode0 = treeselector_.Node("Services", rootNode)
    childNode1 = treeselector_.Node("Photgraphy", childNode0)
    childNode2 = treeselector_.Node("Rental", rootNode)
    childNode3 = treeselector_.Node("LeftTibia", childNode2)
    childNode4 = treeselector_.Node("LeftFoot", childNode3)
    childNode5 = treeselector_.Node("LeftToe", childNode4)

    # print rootNode
    product_services_model = treeselector_.SceneGraphModel(rootNode)
    tree_selector.productStructView.setModel(product_services_model)
    tree = tree_selector.productStructView

    if mode == 'edit':

        tree_selector.pushButton_18.clicked.connect(partial(del_node, tree))
        tree_selector.pushButton_16.clicked.connect(partial(add_node_item, tree))
        tree_selector.pushButton_17.clicked.connect(partial(add_node_subitem, tree))

        unit_price = tree_selector.doubleSpinBox_7.value()

        tree_selector.btnEdit1.clicked.connect(partial(edit_node,tree,
                                                       unit_price, 1,
                                                       tree_selector,
                                                       product_services_model))

        tree_selector.btnEdit2.clicked.connect(partial(edit_node, tree,
                                                       unit_price, 2,
                                                       tree_selector,
                                                       product_services_model))

    print "\n<<Successfully Added Model to Tree selector>>\n"

    return tree_selector


def add_to_db_v2(close_win=None, d=None, table=None):
    """Updated version of adding to the db using ars.

    The point of refactoring function was to provide a was for all views and
    dialogs to update database using common args.
    """
    print "-" * 80
    print "START OF METHOD ADD TO DB V2:"

    connection = sqlite3.connect(DB)
    cursor = connection.cursor()
    # tup = [item for item in d if "NULL" in d[item]]
    # print "check if there is a null: ", d[tup[0]]
    if table == "orders":
        cursor.execute("""INSERT INTO {tn} ("key") VALUES (NULL)""".
                       format(tn=table))
    else:
        cursor.execute("""INSERT INTO {tn} ("key") VALUES (NULL)""".
                       format(tn=table))
    print "Created new row for table"

    cursor.execute("""SELECT DISTINCT "key" FROM {tab}""".format(tab=table))
    result_key = cursor.fetchall()

    last = max(result_key)
    last = last[0]
    print "last key: ", last, "from table:", table
    print "should insert in row last+1: ", last + 1

    if (table and d) is not None:
        print "DATA TO ADD: ", d
        print "\n"
        print "METHOD 'add_to_db_v2' ADDING TO:"
        print "Table:", table
        print "\n"

        for k in d:

            print "Column: ", k
            print "Data:", d[k]

            if d[k] == "NULL":
                print "Truth test:", d[k] == "NULL"
                print "values is:", d[k], "(supposed to be NULL)"
                continue

            else:

                d[k] = "\"" + str(d[k]) + "\""

                print "normal value: ", d[k]
                cursor.execute("""
                    UPDATE {tn} SET "{col}" = {data} WHERE key = {r}""".
                               format(tn=table, col=k, data=d[k], r=last))
                print "Added in col:", k, "value :", d[k]

            print "\n"

    else:
        print "Did not create", table
    print "last row id:", cursor.lastrowid
    connection.commit()
    model.select()
    connection.close()
    print "Finished add_to_db(args)"
    close_win.close()

    print "ENDOF OF METHOD ADD TO DB V2:"
    print "-" * 80


def print_test():
    print "Button connected correctly"


def update_db(close_win=None, d=None, table=None, row=None):
    """Updated version of update to the db using ars.

    The point of refactoring function was to provide a was for all views and
    dialogs to update database using common args.
    """
    print "\nEntered update_db(): ----------------------"
    connection = sqlite3.connect(DB)
    cursor = connection.cursor()

    if (table and row and d) is not None:
        print "\n"
        print "dictionary: ", d
        print "\n"

        print "table:", table
        print "row: ", row
        for k in d:
            print "\n"

            print "col:", k, "type:", type(k)
            print "data:", d[k], "type:", type(d[k])

            d[k] = "'" + d[k] + "'"

            print "\n"

            cursor.execute("""UPDATE {tn} SET "{col}" = {data} WHERE key = {r}"""
                           .format(tn=table, col=k, data=d[k], r=row))
        connection.commit()
        connection.close()

    else:
        print " need a table, a row number and a dict of columns and data !"

    # cursor.execute("SELECT * FROM client")
    # result_all = cursor.fetchall()
    # print "-" * 80
    # print "-" * 80
    # print("fetchall before cc update:")
    # for r in result_all:
    #     print(r)
    # result_cc = [i[0] for i in result_all]
    # print "result_cc = ", result_cc
    # for i in result_cc:
    #     print "i=", i
    #     count = "00000000" + str(i)  # +"\'"
    #     count = count[-3:]

    #     # count = "'" + count + "'"
    #     print "concatenated:", count

    #     cursor.execute(
    #         """UPDATE client SET 'Code Client' = {s}
    #            WHERE 'key' = {k}""" .format(s=count, k=i))

    # connection.commit()
    # cursor.execute("SELECT * FROM client")
    # result_all = cursor.fetchall()
    # print "-" * 80
    # print "-" * 80
    # print("fetchall after cc update:")
    # for r in result_all:
    #     print(r)

    close_win.close()
    print "\nExited update_db(): ----------------------"


def del_node(tree=None):
    index = tree.selectedIndexes()[0]
    print "index>>", index
    tree.model().removeRow(index.row(), parent=index.parent())


def add_node_item(tree=None):
    index = tree.selectedIndexes()[0]
    print "index at item>>", str(index)
    tree.model().insertRow(index.row() + 1, parent=index.parent())
    # tree.model().insertRows(index.row()+1, 1, parent=index.parent())

    model = tree.model()
    # data = []
    # rowCount = model.rowCount(index)
    # print "rowcount form tree:", rowCount
    # for row in range(model.recursiveRowCount(index)):
    #     data.append([])
    #     for column in range(model.columnCount(index)):
    #         index = model.index(row, column)
    #         data[row].append(str(model.data(index, QtCore.Qt.DisplayRole)))
    # print "DATA EXTRACTED FROM TREE", data
    root_index = model.getNode(index, root=True)
    print "root_index is type():", type(root_index)
    print root_index
    print
    print "root_index children:"
    print "%r" % (root_index._children)
    print
    data2 = model.fetchData(root_index)
    print "data2 is type():", type(data2)
    print data2


def add_node_subitem(tree=None):
    index = tree.selectedIndexes()[0]
    print "index at subitem>>", index

    tree.model().insertRows(0, index.row()+1, parent=index)
    # tree.model().insertRows(index.row()+1, 1, parent=index)

    # index = tree.selectedIndexes()[0]
    # print "index>>", index
    # tree.model().removeRow(index.row(), parent=index.parent())


def edit_node(tree=None, value=None, col=None, window=None, model=None):
    if col == 0:
        price = window.doubleSpinBox_7.value()

    elif col == 1:
        price = window.doubleSpinBox_7.value()

    elif col == 2:
        price = window.doubleSpinBox_8.value()

    print
    print "price fetched from SpinBox>>", price
    print

    index = tree.selectedIndexes()[col]

    print "in function edit_node"
    print " value>>", value
    # index.column() = col
    print "index column>>", index.column()

    tree.model().setData(index, price)
    tree.setVisible(False)
    tree.setVisible(True)


def tabView_sel_changed(tabview=None):
    index = tabview.selectedIndexes()[0]
    print('Row %d is selected' % index.row())
    list_of_indices = tabview.selectedIndexes()
    # print "Lst of indices:", list_of_indices
    list_of_data = []
    for index in sorted(list_of_indices):

        # print "Column ", index.column()
        data = str(tabview.model().data(index).toString())
        # print "Data :", data
        list_of_data.append(data)

    return list_of_data

if __name__ == '__main__':

    app = QtGui.QApplication(sys.argv)
    # app = QtGui.QApplication.instance()
    app.setStyle("plastique")
    # app.setMask()
    # set_style(app)

    if not os.path.exists(DB):
        print "DB not found! Creating new DB."
        make_db()

    myDb = QtSql.QSqlDatabase.addDatabase("QSQLITE")
    myDb.setDatabaseName(DB)

    if not myDb.open():
        print "Unable to create connection!"
        sys.exit(1)

    model = MyModel(tab='client')
    model2 = MyModel(tab='job')
    model3 = MyModel(tab='bill')
    model4 = MyModel(tab='invoice_items')
    # app.setStyleSheet('blender_mod_btns.css')

    window = MyWindow(model, model2, model3)

    QtGui.QSizeGrip(window)

    sys.exit(app.exec_())
